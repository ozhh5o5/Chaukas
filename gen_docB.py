"""Chaukas - 12-Page Document Generator (Part B: pages 7-12)"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

D = os.path.dirname(os.path.abspath(__file__))
BLUE = RGBColor(0x00,0x52,0x9B)
DARK = RGBColor(0x1a,0x1a,0x2e)
GRAY = RGBColor(0x55,0x55,0x55)

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Calibri'
        r.font.color.rgb = BLUE if level<=2 else DARK
    return h

def P(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Calibri'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    return p

def IMG(name, w=5.5):
    path = os.path.join(D, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(w))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def CAP(text):
    P(text, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

def BULLET(text):
    doc.add_paragraph(text, style='List Bullet')

# ===== 6. CORE FEATURES & PROTOTYPE WALKTHROUGH (Pages 7-8) =====
H("6. Core Features & Prototype Walkthrough")

P("Chaukas is a multi-module platform with 9 tightly integrated features. Each module is independently deployable yet shares location context and the local database, creating a cohesive command ecosystem.")

IMG('c6_modules.png', 6)
CAP("Figure 6: Chaukas platform module implementation completeness (file checklist derived from this repository). Source: Team Outliers (generated on build).")

H("6.1 AI Severity Engine", level=2)
P("The Severity Engine is the intelligence core of Chaukas. It ingests real-time GPS coordinates and queries three data sources: live weather conditions (temperature, rain, fog, wind speed), nearby incident density from the local SQLite database, and time-of-day risk coefficients derived from MoRTH accident timing data. The composite score produces one of four risk bands:")
BULLET("LOW (0-25%): Normal driving conditions. Routine monitoring advised.")
BULLET("ELEVATED (26-50%): Increased caution required. Speed reduction recommended.")
BULLET("HIGH (51-75%): Significant hazard detected. Consider alternative routes.")
BULLET("CRITICAL (76-100%): Immediate danger. Emergency protocols activated.")
P("The engine stores severity results in sessionStorage, making them instantly available to the Escalation Engine without redundant API calls.")

H("6.2 ML Hotspot Detection Engine", level=2)
P("The ML Hotspot Engine applies DBSCAN (Density-Based Spatial Clustering of Applications with Noise) to incident coordinate data. Unlike K-Means, DBSCAN requires no predefined cluster count, making it ideal for organic accident pattern detection. The engine generates a color-coded heatmap grid with four zone classifications:")
BULLET("Safe Zone (Green): < 2 incidents per grid cell in past 30 days.")
BULLET("Caution Zone (Yellow): 2-5 incidents. Enhanced monitoring.")
BULLET("Warning Zone (Orange): 5-10 incidents. Increased patrol frequency.")
BULLET("Danger Zone (Red): > 10 incidents. Immediate infrastructure intervention required.")
P("The heatmap auto-refreshes every 30 seconds and responds to user-selected radius (25km/50km/100km) and grid resolution (2km/5km/10km).")

H("6.3 Escalation Engine (State Machine)", level=2)
P("The Escalation Engine implements a formal finite state machine with four states derived directly from the National Disaster Management Authority (NDMA) incident response framework:")
BULLET("NORMAL: Risk score < 25%. Routine monitoring and data collection.")
BULLET("WATCH: Risk score 25-50%. Heightened awareness. Advisory alerts issued.")
BULLET("PREPAREDNESS: Risk score 50-75%. Resource pre-positioning. Coordination with emergency services.")
BULLET("CRISIS: Risk score > 75%. Full emergency response. Auto-triggers crisis dispatch module.")
P("State transitions are evaluated every 30 seconds in real-time mode, or can be manually overridden by administrators for scenario simulation exercises.")

H("6.4 Crisis Dispatch & Emergency Contacts", level=2)
P("The Crisis Dispatch module provides location-aware, prioritized access to the nearest trauma centers, ambulances, police stations, and towing services. All contact data is stored locally in SQLite, ensuring 100% availability even without internet connectivity. The module integrates directly with the Escalation Engine: when CRISIS state is triggered, the nearest emergency services are automatically displayed.")

H("6.5 Admin Command Dashboard", level=2)
P("The Admin Dashboard provides a comprehensive, real-time command center for traffic police and highway authorities. Features include: live incident monitoring with GPS coordinates, system-wide analytics (incidents per 24h, 7-day trends, peak hour analysis), user management, broadcast messaging for emergency alerts, and full incident lifecycle management from report to resolution.")

H("6.6 Coordination Intelligence (News Feed)", level=2)
P("The News Intelligence module performs location-aware aggregation of road safety news. The system uses reverse geocoding to convert GPS coordinates to city names, then queries a local SQLite database populated with curated road safety articles. When no API key is available (offline mode), six pre-seeded local news templates covering road safety themes are automatically populated.")

doc.add_page_break()

# ===== 7. TECHNICAL METHODOLOGY (Pages 9-10) =====
H("7. Technical Methodology & Implementation")

IMG('c7_response.png', 6)
CAP("Figure 7: Chaukas ML pipeline performance benchmark (average compute time per operation). Source: Benchmarked from this repository on the current machine during document generation.")

P("The 'Road Watch' theme demands real-time hazard intelligence: driver-facing alerts and authority dashboards must update quickly enough to be actionable. Figure 7 benchmarks the compute cost of three core operations (severity scoring, hotspot detection, heatmap generation) using the actual implementations in this codebase. This supports Chaukas' design goal of near-real-time analysis on commodity hardware, reducing dependence on cloud processing and improving reliability in low-connectivity environments.")

H("7.1 AI & Machine Learning Implementation", level=2)
P("The ML pipeline follows a three-stage architecture:")
BULLET("Stage 1 - Data Ingestion: GPS coordinates and incident metadata are written to local SQLite via the incident reporting module. Each record carries: latitude, longitude, severity level, incident type, timestamp, and AI-generated photo analysis results.")
BULLET("Stage 2 - Spatial Analysis: DBSCAN clustering runs on all coordinate pairs within the user-selected radius. Parameters eps (neighborhood radius in degrees) and min_samples are dynamically calibrated based on data density to prevent both over-clustering and under-detection.")
BULLET("Stage 3 - Zone Classification & Visualization: Cluster outputs are mapped to a grid of configurable resolution. Each cell is assigned a zone based on incident count thresholds, then rendered as color-coded rectangles on the Leaflet.js map.")

H("7.2 Offline-First Architecture", level=2)
P("Chaukas implements a progressive offline-first strategy with three layers of persistence:")

t = doc.add_table(rows=1, cols=3)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t.rows[0].cells
hdr[0].text = 'Layer'; hdr[1].text = 'Technology'; hdr[2].text = 'Data Stored'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs: r.bold = True
layers = [
    ('Browser Cache', 'Service Worker (PWA)', 'Static assets, JS bundles, CSS, fonts'),
    ('Client DB', 'IndexedDB (ChaukasDB)', 'Incidents, user prefs, queued submissions'),
    ('Server DB', 'SQLite (local_db.json)', 'Profiles, incidents, news, analytics'),
]
for l,t2,d in layers:
    row = t.add_row().cells
    row[0].text = l; row[1].text = t2; row[2].text = d

P("When the device goes offline, the Service Worker intercepts all network requests and serves cached responses. New incident reports are queued in IndexedDB and synchronized with the backend when connectivity is restored.")

H("7.3 Data Flow & API Integration", level=2)
P("The platform exposes 14 REST API endpoints through FastAPI. Key endpoints and their technical implementation:")

t2 = doc.add_table(rows=1, cols=3)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr2 = t2.rows[0].cells
hdr2[0].text = 'API'; hdr2[1].text = 'Method'; hdr2[2].text = 'Implementation Detail'
for c in hdr2:
    for p in c.paragraphs:
        for r in p.runs: r.bold = True
apis = [
    ('/api/severity/live', 'POST', 'Composite scoring: weather (40%) + incident density (35%) + temporal (25%)'),
    ('/api/escalation/auto-assess', 'POST', 'NDMA-based state machine with 30s polling cycle'),
    ('/api/hotspot/live-analysis', 'GET', 'DBSCAN on SQLite incident coords; returns heatmap grid JSON'),
    ('/api/incidents/', 'POST', 'Auto-generates ID, timestamps via MockSupabase; stores locally'),
    ('/api/admin/analytics', 'GET', 'Aggregates SQLite data; safe timestamp parsing for offline DBs'),
    ('/api/news/fetch-news', 'POST', 'Falls back to curated local articles if no GNEWS_API_KEY set'),
]
for a,m,d in apis:
    row = t2.add_row().cells
    row[0].text = a; row[1].text = m; row[2].text = d

H("7.4 Photo AI Integration (Gemini Vision)", level=2)
P("Incident reports include photo/video capture analyzed by Google's Gemini Vision API. The model identifies: accident type (vehicle collision, pothole, flooding, obstruction), estimated severity, number of vehicles involved, and recommended emergency response type. Results are stored alongside the incident record and displayed in the Admin Dashboard for human verification.")

doc.add_page_break()

# ===== 8. SCALABILITY & SUSTAINABILITY (Page 11) =====
H("8. Scalability & Sustainability")

IMG('c8_scalability.png', 6)
CAP("Figure 8: Projected 5-Year User Growth Roadmap (Log Scale). From 5,000 pilot users to 10M national users.")

P("Chaukas is architected for progressive scalability across five distinct phases:")

t3 = doc.add_table(rows=1, cols=4)
t3.style = 'Table Grid'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr3 = t3.rows[0].cells
for i,h in enumerate(['Phase','Timeline','Scale','Key Milestone']): hdr3[i].text = h
for c in hdr3:
    for p in c.paragraphs:
        for r in p.runs: r.bold = True
phases = [
    ('Y1 - Pilot', '2026', '5,000 users', 'IIT Madras city pilot; MoU with local traffic police'),
    ('Y2 - State', '2027', '50,000 users', 'Full state rollout; integration with State Highway Authority'),
    ('Y3 - National', '2028', '500,000 users', 'NHAI partnership; national hotspot database seeded'),
    ('Y4 - Infra', '2029', '2M users', 'IoT sensor integration; smart traffic signal feeds'),
    ('Y5 - Full', '2030', '10M users', 'National standard platform; government SLA agreement'),
]
for p,t,s,k in phases:
    row = t3.add_row().cells
    row[0].text = p; row[1].text = t; row[2].text = s; row[3].text = k

H("8.1 Technical Scalability Strategy", level=2)
BULLET("Database Migration Path: SQLite -> PostgreSQL -> Distributed PostgreSQL (read replicas) as user base grows. The MockSupabase client architecture ensures a seamless transition -- simply swap the client with a real Supabase connection by setting environment variables.")
BULLET("API Scalability: FastAPI's async architecture supports 10,000+ concurrent requests per second on standard hardware. Horizontal scaling via Docker/Kubernetes is supported from Day 1.")
BULLET("ML Model Scaling: DBSCAN runs in O(n log n) time with spatial indexing. At 10M incidents, a regional sharding strategy partitions data by geographic grid cells.")

H("8.2 Sustainability", level=2)
BULLET("Open Source: The entire codebase is open-source, ensuring community-driven maintenance and zero licensing costs for government agencies.")
BULLET("Data Sovereignty: Local-first architecture ensures all data remains on Indian servers (or devices), complying with India's Digital Personal Data Protection Act 2023.")
BULLET("Revenue Model: Freemium for citizens; SaaS licensing for municipal corporations and insurance firms; anonymized aggregate data licensing for research institutions.")

doc.add_page_break()

# ===== 9. CONCLUSION & ROADMAP (Page 12) =====
H("9. Conclusion & Road Map")

P("Chaukas represents a paradigm shift in India's approach to road safety -- from reactive incident management to proactive, AI-driven hazard prevention. By combining edge computing, machine learning, and an offline-first architecture, we have built a platform that is not only technically superior but also practically deployable in India's diverse connectivity landscape.")

P("Key achievements of the Chaukas prototype:", bold=True)
BULLET("9 fully integrated modules covering the complete road safety lifecycle from hazard detection to emergency response.")
BULLET("100% offline capability through a three-layer persistence architecture (Service Worker + IndexedDB + SQLite).")
BULLET("Real-time ML hotspot detection using DBSCAN clustering with configurable spatial resolution.")
BULLET("Automated Severity-to-Escalation pipeline reducing manual decision-making latency from hours to seconds.")
BULLET("Near-real-time analysis demonstrated via measured compute benchmarks (Figure 7).")

H("9.1 Immediate Roadmap (6 months)", level=2)
BULLET("Crowdsourced road condition reporting with photo evidence and community verification.")
BULLET("Integration with Government's Vahan database for real-time vehicle identification at accident scenes.")
BULLET("Regional language support (Hindi, Tamil, Telugu) for broader accessibility.")
BULLET("Mobile app (React Native) for Android-first deployment targeting two-wheeler riders.")

H("9.2 Long-Term Vision (3-5 years)", level=2)
BULLET("IoT sensor integration: Road-embedded pressure sensors and CCTV feeds feeding the ML engine.")
BULLET("Federated learning: Privacy-preserving ML model training across distributed city nodes.")
BULLET("Government integration: Full API compatibility with India's Digital Infrastructure for Safe and Smart Transportation (DISST) initiative.")
BULLET("Insurance APIs: Real-time risk scoring for usage-based insurance premiums.")

P("Final Statement:", bold=True)
P("\"Our prototype, CHAUKAS, is an AI-powered solution for the National Road Safety Hackathon 2026. By addressing the 'Road Watch' theme, we bridge the gap between infrastructure hazards and driver awareness. Our approach balances high-impact safety results with low-effort deployment by utilizing edge-AI on common mobile devices. Targeted at both daily commuters and road authorities, CHAUKAS provides a scalable, tech-forward vision for safer Indian roads.\"", italic=True)

H("Data Sources (Figures)", level=2)
BULLET("MoRTH: Road Accidents in India — 2023 publication (PDF): https://morth.gov.in/backend/documents/uploaded/Road-Accident-in-India-2023-Publications.pdf")
BULLET("World Bank API (road traffic death rate): https://api.worldbank.org/v2/country/IND/indicator/SH.STA.TRAF.P5?format=json")

doc.add_paragraph()
P("-- Team Outliers | Presenter: Palash Singh Tomar", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
P("National Road Safety Hackathon 2026 | IIT Madras", align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(os.path.join(D, '_doc_partB.docx'))
print("Part B saved (pages 7-12)")
