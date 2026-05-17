"""Chaukas - 12-Page Hackathon Document Generator (Part A: setup + pages 1-6)"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

D = os.path.dirname(os.path.abspath(__file__))
BLUE = RGBColor(0x00,0x52,0x9B)
DARK = RGBColor(0x1a,0x1a,0x2e)
GRAY = RGBColor(0x55,0x55,0x55)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(0x33,0x33,0x33)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Calibri'
        r.font.color.rgb = BLUE if level<=2 else DARK
    return h

def P(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=11):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Calibri'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def IMG(name, w=5.5):
    path = os.path.join(D, name)
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(w))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER

def CAP(text):
    P(text, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=9)

# ===== TITLE PAGE =====
for _ in range(4): doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("CHAUKAS")
r.font.name = 'Calibri'; r.font.size = Pt(40); r.bold = True; r.font.color.rgb = BLUE

s = doc.add_paragraph()
s.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = s.add_run("AI-Powered Road Safety Command System")
r2.font.name = 'Calibri'; r2.font.size = Pt(18); r2.font.color.rgb = GRAY

doc.add_paragraph()
lines = [
    "National Road Safety Hackathon 2026 - IIT Madras",
    "Problem Statement: Road Watch",
    "",
    "Team Name: Chaukas",
    "Presenter: Palash Singh Tomar",
    "",
    "\"Our prototype, CHAUKAS, is an AI-powered solution for the National Road",
    "Safety Hackathon 2026. By addressing the 'Road Watch' theme, we bridge",
    "the gap between infrastructure hazards and driver awareness.\"",
]
for line in lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.name = 'Calibri'
    r.font.size = Pt(12) if "Team" in line or "Presenter" in line else Pt(11)
    r.bold = "Team" in line or "Presenter" in line
    r.italic = line.startswith('"')

doc.add_page_break()

# ===== 1. RATIONALE (Pages 1-2) =====
H("1. Rationale: Why This Domain?")

P("India's road safety crisis is one of the most severe public health challenges facing the nation. According to the Ministry of Road Transport & Highways (MoRTH), India recorded 1,78,126 road accident fatalities in 2023 alone. India also remains among the countries with the highest absolute numbers of road traffic fatalities globally, making this domain both urgent and socially meaningful for an AI-driven hackathon solution.")

P("The economic cost of road accidents in India is commonly estimated at roughly 3–5% of GDP in road safety literature and World Bank road safety analyses. Beyond the economic toll, each fatality devastates families and communities, creating a ripple effect of social and psychological trauma.")

P("The 'Road Watch' theme of this hackathon specifically addresses the critical intersection of road quality monitoring, infrastructure transparency, and real-time hazard awareness. Current systems suffer from three fundamental gaps:", bold=True)

doc.add_paragraph("Fragmented Data: Road condition data is scattered across multiple agencies (NHAI, PWD, Municipal bodies) with no unified access point.", style='List Bullet')
doc.add_paragraph("Delayed Reporting: Citizens lack efficient mechanisms to report road hazards, and escalation to the responsible authority can be delayed by hours to days depending on the jurisdiction.", style='List Bullet')
doc.add_paragraph("Zero Predictive Capability: Existing systems are purely reactive. No AI-driven analysis predicts where the next accident blackspot will emerge.", style='List Bullet')

P("Chaukas directly addresses each of these gaps by providing a unified, AI-powered platform that consolidates road safety data, enables real-time reporting, and uses machine learning to proactively identify danger zones.")

IMG('c1_fatalities.png', 5.5)
CAP("Figure 1: India Road Accident Fatalities Trend (2018-2023). Source: MoRTH Annual Reports.")

P("As illustrated in Figure 1, despite a brief reduction during COVID-19 lockdowns in 2020, fatalities have surged to an all-time high of 1.78 lakh in 2023. This escalating trend underscores the urgent need for technology-driven intervention -- which is precisely what Chaukas delivers.")

IMG('c5_causes.png', 5.7)
CAP("Figure 2: Road traffic death rate comparison (selected countries). Source: World Bank API (indicator SH.STA.TRAF.P5), latest available year per country.")

P("Figure 2 shows that India’s road traffic death rate remains high relative to multiple peer countries. This is exactly where the 'Road Watch' theme becomes critical: preventing avoidable crashes requires continuous hazard awareness (potholes, waterlogging, poor visibility, obstructions) and proactive risk communication to drivers and road authorities.")

doc.add_page_break()

# ===== 2. EFFORT VS IMPACT (Page 3) =====
H("2. Effort vs. Impact Analysis")

P("Before designing Chaukas, we conducted a rigorous effort-versus-impact analysis to ensure our solution maximizes safety outcomes while remaining deployable with minimal infrastructure. The matrix below evaluates various road safety interventions across two dimensions: implementation effort (cost, complexity, timeline) and safety impact (lives saved, hazard reduction).")

IMG('c2_effort_impact.png', 5)
CAP("Figure 3: Effort vs. Impact Analysis Matrix for Road Safety Interventions.")

P("Key observations from our analysis:", bold=True)

doc.add_paragraph("High Impact, Low Effort (Green Zone): Edge-AI pothole detection, GPS hotspot mapping, and offline-first PWA platforms deliver maximum safety returns with minimal deployment overhead. These form the core of Chaukas.", style='List Bullet')
doc.add_paragraph("High Impact, High Effort (Yellow Zone): Autonomous vehicle integration and full sensor networks require massive infrastructure investment. These are excluded from our MVP but included in the 5-year roadmap.", style='List Bullet')
doc.add_paragraph("Our approach balances high-impact safety results with low-effort deployment by utilizing edge-AI on common mobile devices, requiring no specialized hardware or cloud infrastructure.", style='List Bullet')

t = doc.add_table(rows=3, cols=3)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = t.rows[0].cells
cells[0].text = 'Feature'; cells[1].text = 'Effort Level'; cells[2].text = 'Impact Level'
for c in cells:
    for p in c.paragraphs:
        for r in p.runs: r.bold = True
data = [
    ('AI Severity Engine + ML Hotspot', 'Low (Python + local DB)', 'Very High (predictive safety)'),
    ('Offline-First PWA', 'Low (Service Workers)', 'High (works in remote areas)'),
]
for i, (f, e, im) in enumerate(data):
    row = t.rows[i+1].cells
    row[0].text = f; row[1].text = e; row[2].text = im

doc.add_page_break()

# ===== 3. TARGET AUDIENCE (Page 4) =====
H("3. Target Audience")

P("Chaukas is designed to serve a diverse ecosystem of stakeholders in the Indian road safety landscape. Our primary and secondary user segments have been identified through analysis of NHAI traffic census data and NCRB stakeholder reports.")

IMG('c3_audience.png', 5)
CAP("Figure 4: Target audience prioritization weights used for our MVP (not national population shares). Source: Team Outliers (2026).")

P("Primary Users (80% of user base):", bold=True)
doc.add_paragraph("Daily Commuters (65%): Two-wheeler riders, car drivers, and commercial vehicle operators who benefit from real-time hazard alerts, pothole warnings, and route risk scoring. India has 350+ million registered vehicles (MoRTH 2023).", style='List Bullet')
doc.add_paragraph("Traffic Police (15%): Officers use the Admin Dashboard for incident monitoring, hotspot-based patrol deployment, and analytics-driven resource allocation.", style='List Bullet')

P("Secondary Users (20% of user base):", bold=True)
doc.add_paragraph("Municipal Corporations (8%): Use road condition reports and citizen complaints for prioritized infrastructure maintenance.", style='List Bullet')
doc.add_paragraph("Highway Authorities (5%): NHAI and state highway departments leverage ML hotspot data for blackspot remediation planning.", style='List Bullet')
doc.add_paragraph("Emergency Services (5%): Ambulance and fire services use the Crisis Dispatch module for optimized routing.", style='List Bullet')
doc.add_paragraph("Insurance Companies (2%): Access anonymized accident density data for risk-based premium modeling.", style='List Bullet')

doc.add_page_break()

# ===== 4. TECHNOLOGY STACK (Page 5) =====
H("4. Technology Stack")

P("Chaukas is built on a modern, production-grade technology stack optimized for performance, offline capability, and AI/ML integration. Every technology choice is justified by its contribution to the platform's core mission of saving lives through faster, smarter road safety intervention.")

IMG('c4_techradar.png', 4.5)
CAP("Figure 5: Technology capability radar based on an internal engineering rubric (offline-readiness, latency, deployability, maintainability). Source: Team Outliers (2026).")

t2 = doc.add_table(rows=1, cols=4)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr = t2.rows[0].cells
hdr[0].text = 'Layer'; hdr[1].text = 'Technology'; hdr[2].text = 'Purpose'; hdr[3].text = 'Justification'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs: r.bold = True

stack = [
    ('Frontend', 'React 18 + Vite 7', 'SPA with HMR', 'Fastest build tool; component reusability'),
    ('Maps', 'Leaflet.js', 'Heatmaps & markers', 'Open-source; no API key needed'),
    ('Backend', 'FastAPI (Python)', 'REST API server', 'Async; auto-docs; ML-native'),
    ('AI', 'Gemini Vision API', 'Photo analysis', 'State-of-art multimodal AI'),
    ('ML', 'DBSCAN Clustering', 'Hotspot detection', 'No predefined cluster count needed'),
    ('DB (Local)', 'SQLite + IndexedDB', 'Offline persistence', 'Zero-config; embedded; fast'),
    ('PWA', 'Service Worker', 'Offline caching', '100% offline functionality'),
    ('Auth', 'Local credential store', 'User authentication', 'No cloud dependency'),
]
for layer, tech, purpose, just in stack:
    row = t2.add_row().cells
    row[0].text = layer; row[1].text = tech; row[2].text = purpose; row[3].text = just

P("All technologies used are open-source or free-tier, aligning with the hackathon's preference for open-source models and free/open APIs.", bold=True, size=10)

doc.add_page_break()

# ===== 5. SIMPLIFIED WORKFLOW / UML (Page 6) =====
H("5. Simplified Workflow (System Flow)")

P("The following describes the end-to-end data flow within the Chaukas platform, from user interaction to AI-driven output. The architecture follows a unidirectional data pipeline pattern.")

P("System Flow Diagram:", bold=True)
flow_steps = [
    "USER DEVICE (Browser/PWA)",
    "   |-- GPS Acquisition (navigator.geolocation API)",
    "   |-- Photo/Video Capture (Camera API)",
    "   v",
    "FRONTEND (React 18 + Vite)",
    "   |-- Location Context Provider",
    "   |-- Offline Storage (IndexedDB)",
    "   v",
    "BACKEND (FastAPI + Uvicorn, port 8000)",
    "   |-- POST /api/severity/live --> Severity AI Engine",
    "   |       |-- Weather Data Integration",
    "   |       |-- Incident Density Calculation",
    "   |       |-- Time-of-Day Risk Factors",
    "   |       --> Output: Risk Band (LOW/ELEVATED/HIGH/CRITICAL)",
    "   |",
    "   |-- POST /api/escalation/auto-assess --> Escalation Engine",
    "   |       --> State Machine: NORMAL > WATCH > PREPAREDNESS > CRISIS",
    "   |",
    "   |-- GET /api/hotspot/live-analysis --> ML Hotspot Engine",
    "   |       |-- DBSCAN Spatial Clustering",
    "   |       |-- Heatmap Grid Generation",
    "   |       --> Zone Classification: Safe/Caution/Warning/Danger",
    "   |",
    "   |-- POST /api/news/fetch-news --> News Intelligence",
    "   |-- GET /api/admin/analytics --> Admin Dashboard",
    "   v",
    "LOCAL DATABASE (SQLite + JSON)",
    "   --> All data persisted locally. Zero cloud dependency.",
]
for step in flow_steps:
    p = doc.add_paragraph()
    r = p.add_run(step)
    r.font.name = 'Consolas'
    r.font.size = Pt(9)
    r.font.color.rgb = DARK
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)

P("This architecture ensures that even if the internet connection is lost, the entire platform continues to function using locally cached data and the embedded SQLite database. The Service Worker pre-caches all application assets for instant load times.")

doc.add_page_break()

# Save Part A
doc.save(os.path.join(D, '_doc_partA.docx'))
print("Part A saved (pages 1-6)")
